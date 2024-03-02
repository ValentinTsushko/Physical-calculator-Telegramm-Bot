from docx import Document
from docx.shared import Pt, RGBColor

def reduce_margins(doc):
    # Получаем первый раздел (секцию) в документе
    section = doc.sections[0];

    # Устанавливаем размеры границ (в дюймах)
    section.left_margin = Pt(15.0);
    section.right_margin = Pt(15.0);
    section.top_margin = Pt(0.10);
    section.bottom_margin = Pt(0.10);

# Создаем новый документ
doc = Document();
reduce_margins(doc);

table = doc.add_table(rows=1, cols=2);
table.columns[0].width = Pt(2000.0);
table.columns[1].width = Pt(900.0);

def MakeFirstCol():
    # Make paragraph of first column
    hdr_cells = table.rows[0].cells;
    run = hdr_cells[0].add_paragraph().add_run('Name Surname');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(20);

    run = hdr_cells[0].add_paragraph().add_run('Job title');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(14);

    run = hdr_cells[0].add_paragraph().add_run('CAREER OBJECTIVE');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(0, 128, 0);
    text_under_paragraph1 = (
    "Это много текста, который будет добавлен под первым параграфом. "
    "Вы можете добавить здесь сколько угодно текста."
    );
    run = hdr_cells[0].add_paragraph().add_run(text_under_paragraph1);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);

    run = hdr_cells[0].add_paragraph().add_run('WORK EXPERIENCE');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(0, 128, 0);

    # How many work ex you have for -> n
    # For ex. C/C++ programmer, Town – Corp name
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run('C/C++ programmer'+', ');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(9);
    run = paragraph.add_run('Town'+' – ');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(9);
    run = paragraph.add_run('Corp name');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(9);

    # Start and end work Time
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run('June 2017'+' – ');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);
    run = paragraph.add_run('September 2017');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);

    # How many ex. you solve for -> list.count
    text_under_paragraph1 = (
    """Это много текста,
    который будет добавлен
    под первым параграфом.
    Вы можете добавить здесь
    сколько угодно текста."""
    );
    text_list = text_under_paragraph1.split('\n')
    IsFirstTime = False
    for data in text_list:
        paragraph = hdr_cells[0].add_paragraph('');

        if(not IsFirstTime):
            styles = doc.styles;
            custom_style = styles.add_style('ListBullet', 1);
            custom_style.font.name = 'Times New Roman';
        # Присваиваем стиль параграфу
        paragraph.style = custom_style;
        run = paragraph.add_run(data);
        run.font.size = Pt(7);
        run.font.color.rgb = RGBColor(80, 80, 80);
        IsFirstTime = True

    # Education
    run = hdr_cells[0].add_paragraph().add_run('EDUCATION');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(0, 128, 0);

    # How many ed you have, for -> n
    # For ex. Odesa I.I. Mechnikov National  University – B.Sc.
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run('Odesa I.I. Mechnikov National  University'+' - ');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(11);
    run = paragraph.add_run('B.Sc.');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(11);

    # Start and end work Time
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run('September 2009'+' – ');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);
    run = paragraph.add_run('June 2013');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);

    # How many ex. you solve for -> list.count
    paragraph = hdr_cells[0].add_paragraph('');
    run = paragraph.add_run('Faculty name.');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);

MakeFirstCol();
# Сохраняем документ
doc.save('example_table.docx');
