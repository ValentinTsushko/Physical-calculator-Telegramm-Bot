import telebot
import re
import os
import json
from docx import Document
from docx.shared import Pt, RGBColor


def load_translation(language):
    file_path = os.path.join('Translation', f'cvMaker_{language}.json')
    with open(file_path, 'r', encoding='utf-8') as file:
        translations = json.load(file)
    return translations

def startCVM(message, bot, current_language):
    translations = load_translation(current_language);
    # Создаем новый документ
    doc = Document();
    reduce_margins(doc);
    table = doc.add_table(rows=1, cols=2);
    table.columns[0].width = Pt(2000.0);
    table.columns[1].width = Pt(900.0);
    bot.send_message(message.from_user.id, translations.get("startCVM_text"));
    bot.register_next_step_handler(message, NameSurname, bot, translations, doc, table);


def reduce_margins(doc):
    # Получаем первый раздел (секцию) в документе
    section = doc.sections[0];

    # Устанавливаем размеры границ (в дюймах)
    section.left_margin = Pt(15.0);
    section.right_margin = Pt(15.0);
    section.top_margin = Pt(0.10);
    section.bottom_margin = Pt(0.10);



def NameSurname(message, bot, translations, doc, table):
    # Make paragraph of first column
    hdr_cells = table.rows[0].cells;
    run = hdr_cells[0].add_paragraph().add_run(message.text);
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(20);
    bot.send_message(message.from_user.id, translations.get("jobTitle_text"));
    bot.register_next_step_handler(message, JobTitle, bot, translations, doc, hdr_cells);

def JobTitle(message, bot, translations, doc, hdr_cells):
    run = hdr_cells[0].add_paragraph().add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(14);
    bot.send_message(message.from_user.id, translations.get("CareerObjective_text"));
    bot.register_next_step_handler(message, CareerObjective, bot, translations, doc, hdr_cells);

def CareerObjective(message, bot, translations, doc, hdr_cells):
    run = hdr_cells[0].add_paragraph().add_run('CAREER OBJECTIVE');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(0, 128, 0);
    run = hdr_cells[0].add_paragraph().add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    bot.send_message(message.from_user.id, translations.get("WorkExperience_text"));
    bot.register_next_step_handler(message, WorkExperience, bot, translations, doc, hdr_cells);

def WorkExperience(message, bot, translations, doc, hdr_cells):
    JobsCount = int(message.text)
    print(JobsCount)
    run = hdr_cells[0].add_paragraph().add_run('WORK EXPERIENCE');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(0, 128, 0);
    try:
        #bot.register_next_step_handler(message, WorkCounts, bot, translations, JobsCount, doc, hdr_cells);
        WorkCounts(message, bot, translations, JobsCount, doc, hdr_cells);
    except Exception as e:
        print(f"Error: {e}");

    # How many work ex you have for -> n
    # For ex. C/C++ programmer, Town – Corp name
def WorkCounts(message, bot, translations, JobsCount, doc, hdr_cells):
    print(JobsCount)
    for i in range(JobsCount):
        print(i)
        try:
            bot.send_message(message.from_user.id, translations.get("WorkTitle_text"));
            bot.register_next_step_handler(message, WorkTitle, bot, translations, doc, hdr_cells);
        except Exception as e:
            print(f"Error: {e}");

    bot.send_message(message.from_user.id, translations.get("Education_text"));
    bot.register_next_step_handler(message, Education, bot, translations, doc, hdr_cells);

def WorkTitle(message, bot, translations, doc, hdr_cells):
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run(message.text+', ');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(9);
    bot.send_message(message.from_user.id, translations.get("WorkTown_text"));
    bot.register_next_step_handler(message, WorkTown, bot, translations, doc, hdr_cells);

def WorkTown(message, bot, translations, doc, hdr_cells):
    run = paragraph.add_run(message.text+' – ');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(9);
    bot.send_message(message.from_user.id, translations.get("CorpName_text"));
    bot.register_next_step_handler(message, CorpName, bot, translations, doc, hdr_cells);

def CorpName(message, bot, translations, doc, hdr_cells):
    run = paragraph.add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(9);
    bot.send_message(message.from_user.id, translations.get("StartWork_text"));
    bot.register_next_step_handler(message, StartWorkTime, bot, translations, doc, hdr_cells);

def StartWorkTime(message, bot, translations, doc, hdr_cells):
    # Start work Time
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run(message.text+' – ');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);
    bot.send_message(message.from_user.id, translations.get("EndWork_text"));
    bot.register_next_step_handler(message, EndWorkTime, bot, translations, doc, hdr_cells);

def EndWorkTime(message, bot, translations, doc, hdr_cells):
    # End work Time
    run = paragraph.add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);
    bot.send_message(message.from_user.id, translations.get("ExpWork_text"));
    bot.register_next_step_handler(message, WorkExperience, bot, translations, doc, hdr_cells);

def WorkExperience(message, bot, translations, doc, hdr_cells):
    # How many ex. you solve for -> list.count
    message_text_list = message.text.split('\n');
    IsSecondTime = False;
    for data in message_text_list:
        paragraph = hdr_cells[0].add_paragraph('');
        if(not IsSecondTime):
            styles = doc.styles;
            custom_style = styles.add_style('ListBullet', 1);
            custom_style.font.name = 'Times New Roman';
        paragraph.style = custom_style;
        run = paragraph.add_run(data);
        run.font.size = Pt(7);
        run.font.color.rgb = RGBColor(80, 80, 80);
        IsSecondTime = True

    # Education
def Education(message, bot, translations, doc, hdr_cells):
    EduCount = int(message.text);
    run = hdr_cells[0].add_paragraph().add_run('EDUCATION');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(0, 128, 0);
    EduCounts(message, bot, translations, EduCount, doc, hdr_cells);

def EduCounts(message, bot, translations, EduCount, doc, hdr_cells):
    for i in range(EduCount):
        bot.send_message(message.from_user.id, translations.get("EduTitle_text"));
        bot.register_next_step_handler(message, EduTitle, bot, translations, doc, hdr_cells);
    bot.send_message(message.from_user.id, translations.get("Doc_text"));
    send_document_to_chat(message, bot, doc);

    # How many ed you have, for -> n
    # For ex. Odesa I.I. Mechnikov National  University – B.Sc.
def EduTitle(message, bot, translations, doc, hdr_cells):
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run(message.text+' - ');
    run.font.name = 'Times New Roman';
    run.bold = True;
    run.font.size = Pt(11);
    bot.send_message(message.from_user.id, translations.get("DegreeEdu_text"));
    bot.register_next_step_handler(message, EduDegree, bot, translations, doc, hdr_cells);

def EduDegree(message, bot, translations, doc, hdr_cells):
    run = paragraph.add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(11);
    bot.send_message(message.from_user.id, translations.get("StartEdu_text"));
    bot.register_next_step_handler(message, StartEduTime, bot, translations, doc, hdr_cells);

    # Start and end work Time
def StartEduTime(message, bot, translations, doc, hdr_cells):
    paragraph = hdr_cells[0].add_paragraph();
    run = paragraph.add_run(message.text+' – ');
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);
    bot.send_message(message.from_user.id, translations.get("EndEdu_text"));
    bot.register_next_step_handler(message, EndEduTime, bot, translations, doc, hdr_cells);

def EndEduTime(message, bot, translations, doc, hdr_cells):
    run = paragraph.add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);
    run.font.color.rgb = RGBColor(80, 80, 80);
    bot.send_message(message.from_user.id, translations.get("FacuName_text"));
    bot.register_next_step_handler(message, FacuName, bot, translations, doc, hdr_cells);

def FacuName(message, bot, translations, doc, hdr_cells):
    paragraph = hdr_cells[0].add_paragraph('');
    run = paragraph.add_run(message.text);
    run.font.name = 'Times New Roman';
    run.font.size = Pt(8);

# Сохраняем документ
#doc.save('example_table.docx');
#bot.send_document(chat_id, document=InputFile(document_buffer, filename=document_name))
def send_document_to_chat(message, bot, doc):
    # Создаем буфер в памяти
    doc_buffer = BytesIO(doc);

    # Отправляем документ в чат
    bot.send_document(chat_id=message.chat.id, document=InputFile(doc_buffer, filename="CV"))

#bot.polling(none_stop=True, interval=0, timeout=60)
